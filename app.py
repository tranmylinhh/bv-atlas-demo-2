import streamlit as st
import google.generativeai as genai
from PIL import Image
import docx
import os
import uuid
import pandas as pd
from datetime import datetime

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(page_title="BV-Atlas Marketing", page_icon="img/favicon.png", layout="wide")

# --- CẤU HÌNH ---
BOT_AVATAR = "logo.jpg"
ADMIN_PASSWORD = "admin"  # Mật khẩu để vào xem báo cáo (Bổ sung 1)

# --- 2. CSS GIAO DIỆN (CẬP NHẬT LINK HOVER) ---
st.markdown("""
<style>
    /* 1. Nền tổng thể: Trắng */
    .stApp { background-color: #FFFFFF; color: #000000; }

    /* 2. ÉP BUỘC MỌI VĂN BẢN PHẢI MÀU ĐEN */
    h1, h2, h3, h4, h5, h6, p, div, span, li, .stMarkdown {
        color: #000000 !important;
    }

    /* === 3. CẤU HÌNH LINK (SỬA LẠI TẠI ĐÂY) === */
    /* Trạng thái bình thường */
    a { 
        color: #005792 !important; /* Xanh Bảo Việt chuẩn */
        font-weight: 700 !important;
        text-decoration: none; 
    }
    
    /* Trạng thái khi di chuột vào (Hover) */
    a:hover { 
        color: #002a4d !important; /* Chuyển sang xanh đen đậm hơn */
        text-decoration: underline !important; /* Hiện gạch chân */
        text-decoration-thickness: 2px !important; /* Gạch chân dày hơn chút */
    }

    /* === 4. SIDEBAR === */
    section[data-testid="stSidebar"] { 
        background-color: #F7F9FB; 
        border-right: 1px solid #E0E0E0; 
    }

    /* Label Sidebar màu Đen */
    [data-testid="stSidebar"] .stMarkdown p, 
    [data-testid="stSidebar"] label p,
    [data-testid="stSidebar"] .stRadio label p {
        color: #000000 !important;
        font-weight: 600 !important;
    }

    /* Ô nhập mật khẩu Sidebar */
    [data-testid="stSidebar"] input[type="password"],
    [data-testid="stSidebar"] input[type="text"] {
        background-color: #E8E8E8 !important;
        color: #000000 !important;
        border: 1px solid #999999 !important;
        caret-color: #000000 !important;
    }
    
    [data-testid="stSidebar"] .stTextInput > div > div > input {
        color: #000000 !important;
    }

    /* === 5. BONG BÓNG CHAT === */
    .stChatMessage { padding: 12px 18px; border-radius: 18px; margin-bottom: 10px; display: flex; }
    
    /* Bot (Trái) */
    .stChatMessage[data-testid="stChatMessage"]:nth-child(odd) {
        background-color: #F2F4F6 !important; border: none; flex-direction: row;
    }
    
    /* User (Phải) */
    .stChatMessage[data-testid="stChatMessage"]:nth-child(even) {
        background-color: #E3F2FD !important; border: 1px solid #BBDEFB; flex-direction: row-reverse; text-align: right;
    }
    .stChatMessage[data-testid="stChatMessage"]:nth-child(even) > div:first-child { margin-left: 10px; margin-right: 0; align-items: flex-end; }

    /* === 6. INPUT CHAT CHÍNH === */
    .stChatInput textarea {
        background-color: #F0F2F5 !important;
        color: #000000 !important;
        border: 2px solid #005792 !important;
        border-radius: 30px;
    }
    .stChatInput button { color: #005792 !important; }
    
    /* Header */
    .header-box { text-align: center; margin-bottom: 20px; background: white; padding: 10px; border-bottom: 1px solid #eee; }
    .header-title { color: #005792 !important; font-size: 26px; font-weight: 800; margin: 0; }
    
    #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# --- 3. KẾT NỐI API KEY ---
if 'GOOGLE_API_KEY' in st.secrets:
    genai.configure(api_key=st.secrets['GOOGLE_API_KEY'])
    model = genai.GenerativeModel('gemini-2.0-flash')
else:
    st.error("⚠️ Chưa nhập API Key.")
    st.stop()

# --- 4. HÀM ĐỌC DỮ LIỆU ---
@st.cache_resource
def load_knowledge_base():
    file_path = "Du_lieu_BV_Atlas.docx"
    if not os.path.exists(file_path): return None
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return '\n'.join(full_text)
    except Exception as e: return f"Lỗi đọc file: {e}"

KNOWLEDGE_TEXT = load_knowledge_base()

# Khởi tạo kho Log (Lưu tạm trong phiên làm việc) bổ sung 2
if "logs" not in st.session_state:
    st.session_state.logs = []

def log_data(question, answer, type="Text"):
    # Tự động đánh giá trạng thái
    status = "✅ Thành công"
    if "chưa có thông tin" in answer or "liên hệ" in answer:
        status = "❌ Thiếu dữ liệu (Cần bổ sung)"
    
    st.session_state.logs.append({
        "Thời gian": datetime.now().strftime("%H:%M %d/%m"),
        "Câu hỏi": question,
        "Câu trả lời": answer,
        "Loại": type,
        "Trạng thái": status
    })

# --- 5. SYSTEM PROMPT (UPDATE QUY TẮC KHÔNG SPAM LINK) ---
current_date = datetime.now().strftime("%d/%m/%Y")

SYSTEM_PROMPT = f"""
VAI TRÒ:
Bạn là BV-Atlas, trợ lý ảo chuyên trách của Ban Marketing - Bảo hiểm Bảo Việt.
Đối tượng giao tiếp: Cán bộ nhân viên (Đồng nghiệp) trong công ty.
Mục tiêu: Hỗ trợ đồng nghiệp tìm kiếm tài liệu, thông tin CTKM nhanh nhất để phục vụ kinh doanh.
Avatar: Logo Bảo Việt.

DỮ LIỆU THỜI GIAN: Hôm nay là {current_date}.

QUY TẮC ỨNG XỬ (ƯU TIÊN CAO NHẤT):

1. KHÔNG LIỆT KÊ HÀNG LOẠT (ANTI-SPAM):
   - Nếu User hỏi chung chung (Ví dụ: "Tìm tài liệu", "Gửi link sản phẩm", "Có những gì?"):
     -> TUYỆT ĐỐI KHÔNG liệt kê danh sách link ra ngay.
     -> HÃY HỎI NGƯỢC LẠI để làm rõ nhu cầu: "Chào bạn! Kho tài liệu của mình có đầy đủ thông tin về An Gia, Tâm Bình, K-Care, Intercare... Bạn đang cần tìm cụ thể cho sản phẩm nào ạ?"
   - CHỈ đưa link khi User đã nhắc đến TÊN SẢN PHẨM cụ thể (Ví dụ: "Tài liệu An Gia").
   - TUYỆT ĐỐI KHÔNG GỢI Ý những tài liệu mà bạn KHÔNG CÓ trong tay. (Ví dụ: Đừng hỏi "Bạn có muốn xem biểu phí không?" nếu bạn biết chắc chắn trong kho không có link biểu phí của sản phẩm đó).
   - Chỉ giới thiệu các tài liệu của các sản phẩm có sẵn trong kho dữ liệu cho user.
   - Trường hợp user hỏi về CTKM, hãy cung cấp thông tin chương trình khuyến mãi đang diễn ra, hoặc chương trình khuyến mãi mà họ yêu cầu, không hỏi lại liên quan đến sản phẩm gì mới cung cấp.

2. LOGIC TRẢ LỜI:
   - Bước 1: Xác nhận yêu cầu.
   - Bước 2: Cung cấp đúng thông tin/link của sản phẩm đó (Không kèm sản phẩm khác).
   - Bước 3: Gợi ý mở rộng liên quan đến chính sản phẩm đó.

3. KHI TRAO ĐỔI VỀ CHƯƠNG TRÌNH KHUYẾN MÃI, (QUAN TRỌNG) CẦN KIỂM TRA THỜI HẠN KHUYẾN MÃI:
   - Chỉ liệt kê CTKM có (Ngày kết thúc >= {current_date}).
   - Nếu user hỏi CTKM có (Ngày kết thúc < Hôm nay) -> ĐÃ HẾT HẠN. -> TUYỆT ĐỐI KHÔNG giới thiệu là "đang chạy"/ "Đang diễn ra".
   - Nếu sản phẩm không có CTKM nào đang chạy -> Trả lời thẳng: "Hiện tại sản phẩm này chưa có CTKM mới, bạn có muốn tìm kiếm chương trình khuyến mãi đang chạy không?".

4. XỬ LÝ KHI BỊ BẮT LỖI (Quan trọng):
   - Nếu User phản hồi "Sai rồi", "Hết hạn rồi", "Sao lại giới thiệu cái cũ", "Thông tin này không đúng":
   - HÃY NHẬN LỖI CHÂN THÀNH & TỰ NHIÊN.
   - Ví dụ: "Ôi mình xin lỗi, mình check sót ngày kết thúc. Cảm ơn bạn đã nhắc nhé! Đúng là chương trình này đã hết hạn từ ngày [Ngày] rồi."
   - Đừng vội lôi "văn mẫu" liên hệ Ms. Linh ra ngay, trừ khi Bot thực sự không có dữ liệu để trả lời tiếp.

5. XỬ LÝ KHI THIẾU THÔNG TIN / USER KHÓ CHỊU:
   - Nếu không tìm thấy mà không thể thay thế bằng một thông tin khác:
     "Thành thật xin lỗi bạn vì sự bất tiện này 😔. Kho dữ liệu của mình chưa có tài liệu này rồi, Ban Marketing vẫn đang cập nhật thêm dữ liệu. Nếu cần gấp, bạn vui lòng nhắn trực tiếp đầu mối Ban Marketing là Ms. TRẦN MỸ LINH (tran.my.linh@baoviet.com.vn) để được hỗ trợ ngay nhé!"

6. PHONG CÁCH, THÁI ĐỘ "ĐỒNG NGHIỆP" (Human Touch):
   - Không trả lời như người máy vô cảm. Hãy dùng ngôn ngữ văn phòng, lịch sự nhưng gần gũi, thân thiện.
   - Xưng hô: "Mình" - "Bạn".
   - Dùng emoji 😊, ☺️, 🥹 để giảm căng thẳng.
"""
# --- KHỞI TẠO SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "type": "text", "content": f"Chào bạn! 👋 Mình là BV-Atlas đây. Bạn cần tìm tài liệu hay check thông tin chương trình khuyến mãi hôm nay?"}
    ]

# Khởi tạo ID cho nút upload (Chìa khóa để fix lỗi đỏ)
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = str(uuid.uuid4())

# --- 6. GIAO DIỆN CHÍNH ---

# === SIDEBAR: CHUYỂN ĐỔI USER / ADMIN ===
with st.sidebar:
    st.image(BOT_AVATAR, width=120)
    
    # Menu chuyển đổi
    app_mode = st.radio("Chế độ xem:", ["👤 Nhân viên Tra cứu", "🔐 Admin Báo cáo"])
    st.markdown("---")

    if app_mode == "👤 Nhân viên Tra cứu":
        st.markdown("### 📸 Tra cứu Ảnh")
        st.info("Upload ảnh Poster/Banner để hỏi thông tin.")
        uploaded_img = st.file_uploader("Chọn ảnh...", type=['jpg', 'png', 'jpeg'], label_visibility="collapsed", key=f"uploader_{st.session_state.get('uploader_key', 'init')}")
        
        img_data = None
        if uploaded_img:
            img_data = Image.open(uploaded_img)
            st.image(img_data, caption="Ảnh xem trước", use_container_width=True)

# === LOGIC MÀN HÌNH CHÍNH ===

if app_mode == "🔐 Admin Báo cáo":
    # === GIAO DIỆN ADMIN ===
    st.title("📊 Báo cáo Tra cứu BV-Atlas")
    
    password = st.text_input("Nhập mật khẩu Admin:", type="password")
    if password == "admin123": # Mật khẩu Demo
        if len(st.session_state.logs) > 0:
            df = pd.DataFrame(st.session_state.logs)
            
            # Metrics
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("Tổng lượt hỏi", len(df))
            col_b.metric("Câu hỏi Thiếu dữ liệu", len(df[df['Trạng thái'].str.contains("Thiếu")]))
            col_c.metric("Tỉ lệ đáp ứng", f"{100 - (len(df[df['Trạng thái'].str.contains("Thiếu")])/len(df)*100):.0f}%")
            
            st.markdown("### 📝 Nhật ký chi tiết")
            st.dataframe(df, use_container_width=True)
            
            # Nút tải về
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Tải báo cáo về máy (Excel/CSV)", csv, "bao_cao_bvatlas.csv", "text/csv")
        else:
            st.info("Chưa có dữ liệu tra cứu nào trong phiên này. Hãy quay lại tab 'Nhân viên' và chat thử vài câu!")
    elif password:
        st.error("Sai mật khẩu!")

else:
    # === GIAO DIỆN CHAT (USER) ===
    
    # Header
    st.markdown(f"""
        <div class="header-container">
            <img src="{BOT_AVATAR}" width="60" style="vertical-align: middle;">
            <div class="header-title">BV-Atlas Marketing</div>
        </div>
    """, unsafe_allow_html=True)

    if KNOWLEDGE_TEXT is None:
        st.warning("⚠️ Chưa tìm thấy file dữ liệu.")

    # Lịch sử Chat
    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "assistant", "type": "text", "content": f"Chào bạn! 👋 Mình là BV-Atlas. Bạn cần tìm tài liệu hay check khuyến mãi gì hôm nay?"}]
    if "uploader_key" not in st.session_state:
        st.session_state.uploader_key = str(uuid.uuid4())

    for msg in st.session_state.messages:
        if msg["role"] == "assistant":
            with st.chat_message(msg["role"], avatar=BOT_AVATAR): st.markdown(msg["content"])
        else:
            with st.chat_message(msg["role"], avatar="👤"):
                if msg.get("type") == "image": st.image(msg["content"], width=200)
                else: st.markdown(msg["content"])

    # Input Chat
    if prompt := st.chat_input("Nhập câu hỏi..."):
        # Xử lý User
        if img_data:
            st.session_state.messages.append({"role": "user", "type": "image", "content": img_data})
            with st.chat_message("user", avatar="👤"): st.image(img_data, width=200)
            
        st.session_state.messages.append({"role": "user", "type": "text", "content": prompt})
        with st.chat_message("user", avatar="👤"): st.markdown(prompt)

        # Bot trả lời & Ghi Log
        with st.chat_message("assistant", avatar=BOT_AVATAR):
            with st.spinner("..."):
                try:
                    history_text = ""
                    for msg in st.session_state.messages[-5:]:
                        if msg.get("type") == "text":
                            role_name = "User" if msg["role"] == "user" else "BV-Atlas"
                            history_text += f"{role_name}: {msg['content']}\n"

                    final_prompt = [
                        f"{SYSTEM_PROMPT}\n",
                        f"=== DỮ LIỆU NỘI BỘ ===\n{KNOWLEDGE_TEXT}\n",
                        f"=== LỊCH SỬ CHAT ===\n{history_text}\n",
                        f"CÂU HỎI USER: {prompt}"
                    ]
                    
                    if img_data:
                        final_prompt.append("User gửi ảnh. Hãy phân tích.")
                        final_prompt.append(img_data)
                    
                    response = model.generate_content(final_prompt)
                    
                    st.markdown(response.text)
                    st.session_state.messages.append({"role": "assistant", "type": "text", "content": response.text})
                    
                    # --- GHI LOG (Quan trọng) ---
                    log_data(prompt, response.text, "Image" if img_data else "Text")
                    
                    st.session_state.uploader_key = str(uuid.uuid4())
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Lỗi: {e}")
